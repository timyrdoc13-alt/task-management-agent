"""Export research markdown to DOCX (readable in Kaiten) and plain-text card summaries."""

from __future__ import annotations

import base64
import io
import os
import re
import urllib.error
import urllib.request
import zlib
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt


def _strip_md_inline(text: str) -> str:
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1", text)
    text = re.sub(r"https?://\S+", "", text)
    text = re.sub(r"\s+", " ", text).strip(" ()")
    return text.strip()


def _kroki_encode(diagram: str) -> str:
    compressed = zlib.compress(diagram.encode("utf-8"), 9)
    return base64.urlsafe_b64encode(compressed).decode("ascii")


def _env_get(key: str, default: str = "") -> str:
    try:
        from kaiten_api import ENV  # noqa: WPS433

        return ENV.get(key, os.environ.get(key, default))
    except Exception:
        return os.environ.get(key, default)


def render_mermaid_png(mermaid_code: str) -> bytes | None:
    """Render mermaid diagram via Kroki (public) or self-hosted KROKI_BASE_URL."""
    if _env_get("RESEARCH_CHARTS_ENABLED", "true").lower() in {"0", "false", "no"}:
        return None
    base = _env_get("KROKI_BASE_URL", "https://kroki.io").rstrip("/")
    encoded = _kroki_encode(mermaid_code.strip())
    url = f"{base}/mermaid/png/{encoded}"
    req = urllib.request.Request(url, headers={"User-Agent": "kaiten-agent/1.0"})
    try:
        with urllib.request.urlopen(req, timeout=45) as resp:
            data = resp.read()
            return data if len(data) > 100 else None
    except (urllib.error.URLError, TimeoutError, OSError):
        return None


def markdown_to_docx(markdown: str, out_path: Path) -> Path:
    """Convert research markdown to Word; embed mermaid diagrams as PNG when possible."""
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    lines = markdown.splitlines()
    i = 0
    pending_caption: str | None = None

    while i < len(lines):
        line = lines[i].rstrip()
        if not line.strip():
            i += 1
            continue

        if line.startswith("```"):
            lang = line[3:].strip().lower() or "text"
            buf: list[str] = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                buf.append(lines[i])
                i += 1
            if i < len(lines):
                i += 1
            code = "\n".join(buf)
            if lang == "mermaid":
                png = render_mermaid_png(code)
                if pending_caption:
                    doc.add_paragraph(pending_caption)
                    pending_caption = None
                if png:
                    doc.add_picture(io.BytesIO(png), width=Inches(5.8))
                else:
                    doc.add_paragraph("[Диаграмма mermaid — открой report.md или включи Kroki]")
                    doc.add_paragraph(code[:2000])
            else:
                doc.add_paragraph(code[:3000])
            continue

        if line.startswith("# "):
            doc.add_heading(_strip_md_inline(line[2:]), level=0)
            i += 1
            continue
        if line.startswith("## "):
            doc.add_heading(_strip_md_inline(line[3:]), level=1)
            i += 1
            continue
        if line.startswith("### "):
            doc.add_heading(_strip_md_inline(line[4:]), level=2)
            i += 1
            continue

        if line.lower().startswith("рис.") or line.lower().startswith("fig."):
            pending_caption = _strip_md_inline(line)
            i += 1
            continue

        m_num = re.match(r"^(\d+)\.\s+(.+)$", line)
        if m_num:
            p = doc.add_paragraph(style="List Number")
            p.add_run(_strip_md_inline(m_num.group(2)))
            i += 1
            continue

        if re.match(r"^[-*]\s+", line):
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(_strip_md_inline(re.sub(r"^[-*]\s+", "", line)))
            i += 1
            continue

        if "|" in line and line.strip().startswith("|"):
            table_rows: list[list[str]] = []
            while i < len(lines) and "|" in lines[i]:
                row = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                if row and not all(re.match(r"^[-:]+$", c) for c in row):
                    table_rows.append([_strip_md_inline(c) for c in row])
                i += 1
            if table_rows:
                tbl = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
                tbl.style = "Table Grid"
                for ri, row in enumerate(table_rows):
                    for ci, cell in enumerate(row[: len(table_rows[0])]):
                        tbl.rows[ri].cells[ci].text = cell[:500]
            continue

        doc.add_paragraph(_strip_md_inline(line))
        i += 1

    out_path = Path(out_path).expanduser().resolve()
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path


def export_report_files(md_path: Path) -> dict[str, str]:
    """Write report.docx next to report.md; return paths."""
    md_path = Path(md_path).expanduser().resolve()
    text = md_path.read_text(encoding="utf-8")
    docx_path = md_path.with_name("report.docx")
    markdown_to_docx(text, docx_path)
    return {"markdown": str(md_path), "docx": str(docx_path)}


def extract_section(markdown: str, heading: str) -> str:
    """Extract body of ## heading until next ## or #."""
    pattern = rf"(?m)^##\s+{re.escape(heading)}\s*\n(.*?)(?=^##\s|\Z)"
    m = re.search(pattern, markdown, re.S)
    return m.group(1).strip() if m else ""


def kaiten_card_description(markdown: str, topic: str) -> str:
    """Plain-text description for Kaiten (no markdown — poor renderer)."""
    tldr = extract_section(markdown, "TL;DR") or extract_section(markdown, "Кратко")
    steps = extract_section(markdown, "Дальнейшие шаги") or extract_section(markdown, "Следующие шаги")

    lines = [f"Ресёрч (ИИ/ЦТ): {topic[:200]}", ""]
    if tldr:
        lines.append("КРАТКО")
        for raw in tldr.splitlines():
            raw = raw.strip()
            if not raw:
                continue
            raw = re.sub(r"^[-*]\s+", "", raw)
            raw = re.sub(r"^\d+\.\s+", "", raw)
            raw = _strip_md_inline(raw)
            if raw:
                lines.append(f"• {raw[:500]}")
        lines.append("")
    else:
        body = re.sub(r"^#.+$", "", markdown, flags=re.M).strip()
        snippet = "\n".join(body.splitlines()[:6])
        if snippet:
            lines.append(snippet[:800])
            lines.append("")

    if steps:
        lines.append("ДАЛЬШЕ")
        for raw in steps.splitlines()[:5]:
            raw = raw.strip()
            if raw:
                step = _strip_md_inline(re.sub(r"^[-*]\s+", "", raw))[:300]
                lines.append(f"• {step}")
        lines.append("")

    lines.append("Полный отчёт с диаграммами — во вложении report.docx.")
    text = "\n".join(lines)
    cap = int(_env_get("KAITEN_DESC_MAX_CHARS", "8000"))
    return text[:cap]
